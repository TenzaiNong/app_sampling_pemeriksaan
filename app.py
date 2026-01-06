import math
import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import csv
import chardet  # Perlu di-install: pip install chardet

# --- IMPORT MODUL SENDIRI ---
import calculations as calc
import selections as sel

# --- FUNGSI HELPER ---
def detect_encoding(uploaded_file):
    """
    Deteksi encoding file CSV.
    Coba: utf-8, latin-1, cp1252, iso-8859-1, windows-1252
    """
    try:
        uploaded_file.seek(0)
        raw_data = uploaded_file.read()
        
        # Gunakan chardet untuk deteksi encoding
        result = chardet.detect(raw_data)
        encoding = result.get('encoding', 'utf-8')
        
        uploaded_file.seek(0)
        return encoding if encoding else 'utf-8'
    except Exception as e:
        # Fallback ke encoding umum Windows
        uploaded_file.seek(0)
        return 'latin-1'

def detect_csv_delimiter(uploaded_file, sample_size=5000):
    """
    Deteksi delimiter CSV secara otomatis.
    Coba: ',', ';', '\t', '|'
    """
    encoding = detect_encoding(uploaded_file)
    
    try:
        # Baca sample dari file dengan encoding yang terdeteksi
        uploaded_file.seek(0)
        sample = uploaded_file.read(sample_size).decode(encoding, errors='ignore')
        
        # Coba deteksi dengan csv.Sniffer
        delimiter = csv.Sniffer().sniff(sample).delimiter
        uploaded_file.seek(0)
        return delimiter, encoding
    except Exception as e:
        # Jika gagal, coba delimiter umum
        delimiters = [';', ',', '\t', '|']
        
        for delim in delimiters:
            try:
                uploaded_file.seek(0)
                test_df = pd.read_csv(uploaded_file, sep=delim, nrows=5, encoding=encoding)
                if len(test_df.columns) > 1:
                    uploaded_file.seek(0)
                    return delim, encoding
            except:
                continue
        
        uploaded_file.seek(0)
        return ',', encoding  # Default ke comma
    
def convert_rupiah_to_numeric(df):
    """
    Konversi kolom dengan format Rupiah (1.234.567,50) ke numerik.
    Format: Pemisah ribuan '.' dan desimal ','.
    """
    for col in df.columns:
        try:
            # Cek jika kolom berisi tanda Rupiah
            if df[col].dtype == 'object':  # Kolom string
                # Ambil sample untuk deteksi
                sample = df[col].astype(str).iloc[0]
                
                # Deteksi jika format Rupiah (ada . dan ,)
                if '.' in sample and ',' in sample:
                    # Contoh: "1.234.567,50" -> "1234567.50"
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace('.', '', regex=False)  # Hapus pemisah ribuan
                               .str.replace(',', '.', regex=False)  # Ubah desimal , menjadi .
                               .apply(pd.to_numeric, errors='coerce'))
                    
                    print(f"✅ Konversi kolom '{col}' dari format Rupiah ke numerik")
                
                elif ',' in sample and '.' not in sample:
                    # Alternatif: hanya ada , sebagai desimal
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace(',', '.', regex=False)
                               .apply(pd.to_numeric, errors='coerce'))
                    
                    print(f"✅ Konversi kolom '{col}' (desimal koma) ke numerik")
                    
        except Exception as e:
            pass  # Skip kolom yang tidak bisa dikonversi
    
    return df

#KODE untuk buat Laporan dalam bentuk excel
def generate_laporan_xlsx(df_original, sampled_df, metode_sampling, teknik, 
                          confidence, sst, value_col, n_final):
    """
    Generate laporan hasil sampling dalam format Excel (.xlsx)
    dengan multiple sheets: Ringkasan, Metode, Detail Sampel
    """
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    buff = BytesIO()
    
    with pd.ExcelWriter(buff, engine='openpyxl') as writer:
        # === SHEET 1: RINGKASAN ===
        ringkasan_data = {
            'Keterangan': [
                'Total Populasi (N)',
                'Total Nilai Buku (Rp)',
                'Jumlah Sampel (n)',
                'Persentase Sampel (%)',
                'Metode Sampling',
                'Teknik Pemilihan',
                'Confidence Level (%)',
                'Salah Saji Tertoleransi (Rp)',
                'Tanggal Generate Laporan'
            ],
            'Nilai': [
                len(df_original),
                f"Rp {df_original[value_col].sum():,.2f}",
                n_final,
                f"{(n_final / len(df_original) * 100):.2f}%",
                metode_sampling,
                teknik,
                confidence,
                f"Rp {sst:,.2f}",
                pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S")
            ]
        }
        df_ringkasan = pd.DataFrame(ringkasan_data)
        df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
        
        # === SHEET 2: DETAIL SAMPEL ===
        sampled_df.to_excel(writer, sheet_name='Detail Sampel', index=False)
        
        # === SHEET 3: STATISTIK POPULASI ===
        statistik_data = {
            'Metrik': [
                'Mean',
                'Median',
                'Std Dev',
                'Min',
                'Max',
                'Q1 (25%)',
                'Q3 (75%)'
            ],
            'Populasi': [
                f"Rp {df_original[value_col].mean():,.2f}",
                f"Rp {df_original[value_col].median():,.2f}",
                f"Rp {df_original[value_col].std():,.2f}",
                f"Rp {df_original[value_col].min():,.2f}",
                f"Rp {df_original[value_col].max():,.2f}",
                f"Rp {df_original[value_col].quantile(0.25):,.2f}",
                f"Rp {df_original[value_col].quantile(0.75):,.2f}"
            ],
            'Sampel': [
                f"Rp {sampled_df[value_col].mean():,.2f}",
                f"Rp {sampled_df[value_col].median():,.2f}",
                f"Rp {sampled_df[value_col].std():,.2f}",
                f"Rp {sampled_df[value_col].min():,.2f}",
                f"Rp {sampled_df[value_col].max():,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.25):,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.75):,.2f}"
            ]
        }
        df_statistik = pd.DataFrame(statistik_data)
        df_statistik.to_excel(writer, sheet_name='Statistik', index=False)
        
        # === FORMATTING ===
        workbook = writer.book
        
        # Format Ringkasan
        ws_ringkasan = workbook['Ringkasan']
        fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        font_header = Font(bold=True, color="FFFFFF")
        
        for cell in ws_ringkasan[1]:
            cell.fill = fill_header
            cell.font = font_header
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Auto width
        for column in ws_ringkasan.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws_ringkasan.column_dimensions[column_letter].width = adjusted_width
        
        # Format Detail Sampel
        ws_detail = workbook['Detail Sampel']
        for cell in ws_detail[1]:
            cell.fill = fill_header
            cell.font = font_header
        
        # Format Statistik
        ws_statistik = workbook['Statistik']
        for cell in ws_statistik[1]:
            cell.fill = fill_header
            cell.font = font_header
    
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK GENERATE LAPORAN DOCX ---
def generate_laporan_docx(df_original, sampled_df, metode_sampling, teknik,
                          confidence, sst, value_col, n_final, max_rows=300):
    """
    Buat laporan .docx:
    - Ringkasan metadata
    - Statistik populasi vs sampel
    - Potongan Detail Sampel (maks max_rows)
    Mengembalikan BytesIO yang siap di-download.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Judul
    h = doc.add_heading('Laporan Sampling Pemeriksaan Keuangan', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata / Ringkasan singkat
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tanggal: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S"))
    doc.add_paragraph(f"Metode Sampling: {metode_sampling}")
    doc.add_paragraph(f"Teknik Pemilihan: {teknik}")
    doc.add_paragraph(f"Confidence Level: {confidence}%")
    doc.add_paragraph(f"Salah Saji Tertoleransi (SST): Rp {sst:,.2f}")
    doc.add_paragraph(f"Jumlah Populasi (N): {len(df_original)}")
    if value_col:
        try:
            total_nilai = df_original[value_col].sum()
            doc.add_paragraph(f"Total Nilai Buku: Rp {total_nilai:,.2f}")
        except Exception:
            doc.add_paragraph("Total Nilai Buku: -")
    doc.add_paragraph(f"Jumlah Sampel (n): {n_final}")
    doc.add_paragraph()

    # Statistik: buat tabel metric vs populasi vs sampel
    metrics = ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Q1 (25%)', 'Q3 (75%)']
    pop_values = []
    samp_values = []
    if value_col and value_col in df_original.columns and value_col in sampled_df.columns:
        pop = df_original[value_col].dropna().astype(float)
        samp = sampled_df[value_col].dropna().astype(float)
        pop_values = [
            f"Rp {pop.mean():,.2f}",
            f"Rp {pop.median():,.2f}",
            f"Rp {pop.std():,.2f}",
            f"Rp {pop.min():,.2f}",
            f"Rp {pop.max():,.2f}",
            f"Rp {pop.quantile(0.25):,.2f}",
            f"Rp {pop.quantile(0.75):,.2f}"
        ]
        samp_values = [
            f"Rp {samp.mean():,.2f}" if not samp.empty else "-",
            f"Rp {samp.median():,.2f}" if not samp.empty else "-",
            f"Rp {samp.std():,.2f}" if not samp.empty else "-",
            f"Rp {samp.min():,.2f}" if not samp.empty else "-",
            f"Rp {samp.max():,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.25):,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.75):,.2f}" if not samp.empty else "-"
        ]
    else:
        pop_values = ['-'] * len(metrics)
        samp_values = ['-'] * len(metrics)

    doc.add_heading('Statistik Populasi vs Sampel', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = 'Metrik'
    hdr_cells[1].text = 'Populasi'
    hdr_cells[2].text = 'Sampel'
    for i, m in enumerate(metrics):
        row_cells = tbl.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = pop_values[i]
        row_cells[2].text = samp_values[i]

    doc.add_paragraph()

    # Detail Sampel (batasi baris)
    doc.add_heading(f'Detail Sampel (menampilkan max {max_rows} baris)', level=2)
    if sampled_df.empty:
        doc.add_paragraph("Tidak ada sampel terpilih.")
    else:
        # batasi kolom agar tabel Word tidak terlalu lebar
        max_cols = 10
        cols = list(sampled_df.columns)[:max_cols]
        sub = sampled_df[cols].head(max_rows).fillna('')

        table = doc.add_table(rows=1, cols=len(cols))
        # header
        for j, c in enumerate(cols):
            table.rows[0].cells[j].text = str(c)
        # rows
        for _, r in sub.iterrows():
            cells = table.add_row().cells
            for j, c in enumerate(cols):
                val = r[c]
                # convert to str safely
                cells[j].text = str(val)

        if len(sampled_df.columns) > max_cols:
            doc.add_paragraph(f"... (kolom terpotong; dokumen menampilkan hanya {max_cols} kolom pertama).")

        if len(sampled_df) > max_rows:
            doc.add_paragraph(f"... (baris terpotong; dokumen menampilkan hanya {max_rows} baris pertama).")

    # Footer / catatan
    doc.add_paragraph()
    doc.add_paragraph("Catatan: File ini dihasilkan otomatis. Untuk tabel lengkap gunakan format .xlsx.")

    # Simpan ke BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Callback yang dipanggil saat tombol Generate .docx diklik
def on_generate_docx():
    try:
        sampled_df = st.session_state.get('sampled_df', pd.DataFrame())
        if sampled_df is None or sampled_df.empty:
            st.warning("Belum ada sampel. Klik 'Generate Sampel' terlebih dahulu.")
            return

        # panggil fungsi pembuat docx yang Anda miliki; pastikan mengembalikan BytesIO
        buf = generate_laporan_docx(
            df_original=df,             # pastikan `df` adalah dataframe populasi yang ada di scope
            sampled_df=sampled_df,
            metode_sampling=st.session_state.get('sampled_params', {}).get('metode_sampling', metode_sampling),
            teknik=st.session_state.get('sampled_params', {}).get('teknik', teknik),
            confidence=confidence,
            sst=sst,
            value_col=value_col,
            n_final=n_final
        )
        # simpan bytes ke session_state agar tidak dihitung ulang
        st.session_state['report_docx_bytes'] = buf.getvalue()
        st.success("Laporan .docx berhasil dibuat. Tombol download muncul di bawah.")
    except Exception as exc:
        st.error(f"Gagal membuat laporan .docx: {exc}")
        # untuk debugging, tunjukkan traceback singkat (opsional)
        import traceback
        st.text(traceback.format_exc())


st.set_page_config(page_title="Dashboard Sampling Audit", layout="wide")
st.title("🕵️ Dashboard Uji Petik Pemeriksaan Keuangan")
st.markdown("---")

# 1. UPLOAD
st.sidebar.header("1. Upload Data")
uploaded_file = st.sidebar.file_uploader("Upload Tabel Data (Excel/CSV/Parquet)",
                                         type=['xlsx', 'csv', 'parquet'])

if uploaded_file is not None:
    # Load Data
    if uploaded_file.name.endswith('.csv'):
        try:
            # Deteksi delimiter dan encoding
            delimiter, encoding = detect_csv_delimiter(uploaded_file)
            df = pd.read_csv(uploaded_file, sep=delimiter, encoding=encoding)
            
            # 🔥 KONVERSI FORMAT RUPIAH KE NUMERIK
            df = convert_rupiah_to_numeric(df)
            
            st.info(f"📌 Delimiter: **'{delimiter}'** | Encoding: **{encoding}**")
        except Exception as e:
            st.error(f"❌ Gagal membaca file CSV: {e}")
            st.stop()
    
    elif uploaded_file.name.endswith('.parquet'):
        try:
            df = pd.read_parquet(uploaded_file)
            st.info("✅ File Parquet berhasil dibaca")
        except Exception as e:
            st.error(f"❌ Gagal membaca file Parquet: {e}")
            st.stop()
    
    else:  # Excel (.xlsx)
        try:
            df = pd.read_excel(uploaded_file)
            st.info("✅ File Excel berhasil dibaca")
        except Exception as e:
            st.error(f"❌ Gagal membaca file Excel: {e}")
            st.stop()

    st.success(f"Data dimuat: {len(df)} baris.")
    
    # untuk menampilkan preview data
    # st.write("📋 Preview Data:")
    # st.dataframe(df.head())
    # st.write("Tipe Data Kolom:")
    #st.write(df.dtypes)

    # Konversi kolom dengan format Rupiah (1.234.567,50) ke numerik.
    # Format: Pemisah ribuan '.' dan desimal ','.
    for col in df.columns:
        try:
            # Cek jika kolom berisi tanda Rupiah
            if df[col].dtype == 'object':  # Kolom string
                # Ambil sample untuk deteksi
                sample = df[col].astype(str).iloc[0]
                
                # Deteksi jika format Rupiah (ada . dan ,)
                if '.' in sample and ',' in sample:
                    # Contoh: "1.234.567,50" -> "1234567.50"
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace('.', '', regex=False)  # Hapus pemisah ribuan
                               .str.replace(',', '.', regex=False)  # Ubah desimal , menjadi .
                               .apply(pd.to_numeric, errors='coerce'))
                    
                    print(f"✅ Konversi kolom '{col}' dari format Rupiah ke numerik")
                
                elif ',' in sample and '.' not in sample:
                    # Alternatif: hanya ada , sebagai desimal
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace(',', '.', regex=False)
                               .apply(pd.to_numeric, errors='coerce'))
                    
                    print(f"✅ Konversi kolom '{col}' (desimal koma) ke numerik")
                    
        except Exception as e:
            pass  # Skip kolom yang tidak bisa dikonversi
    
    # 2. PEMETAAN
    col1, col2 = st.columns(2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    all_cols = df.columns.tolist()

    with col1:
        id_col = st.selectbox("Kolom ID Sampel", all_cols)
    with col2:
        value_col = st.selectbox("Kolom Nilai Rupiah", numeric_cols)

    total_nilai_buku = df[value_col].sum()
    st.info(f"💰 Total Nilai Buku: Rp {total_nilai_buku:,.2f}")

    # 3. METODE HITUNG JUMLAH SAMPEL
    st.markdown("---")
    st.subheader("Metode Penentuan Jumlah Sampel")

    metode_sampling = st.selectbox("Pilih Metode", [
        "Monetary Unit Sampling (MUS)", "Unstratified Mean Per Unit (MPU)",
        "Stratified Mean Per Unit (MPU)", "Difference/Ratio Estimation"
    ])

    # Parameter Input UI
    col_in1, col_in2, col_in3 = st.columns(3)
    with col_in1:
        confidence = st.selectbox("Confidence Level", [90, 95, 99], index=1)
    with col_in2:
        sst = st.number_input("Salah Saji Tertoleransi (SST)",
                              value=float(total_nilai_buku) * 0.05)

    n_res = 0
    error_msg = None

    # --- PANGGIL FUNGSI DARI MODUL CALCULATIONS ---
    if metode_sampling == "Monetary Unit Sampling (MUS)":
        with col_in3:
            dss = st.number_input("Dugaan Salah Saji (DSS)", value=0.0)
        n_res, error_msg = calc.calculate_mus(total_nilai_buku, confidence,
                                              sst, dss)

    elif metode_sampling == "Unstratified Mean Per Unit (MPU)":
        with col_in3:
            sd = st.number_input("Estimasi Standar Deviasi", value=100000.0)
        n_res, error_msg = calc.calculate_mpu_unstratified(
            len(df), confidence, sst, sd)

    elif metode_sampling == "Stratified Mean Per Unit (MPU)":
        st.info("ℹ️ Stratifikasi otomatis menggunakan metode Kuantil (Membagi populasi sama rata).")
        
        # 1. KONSEP PILIHAN KUARTIL
        # Default audit biasanya 4 (Kuartil), tapi kita beri fleksibilitas
        n_bins = st.slider("Bagi populasi menjadi berapa bagian? (Default: 4 - Kuartil)", min_value=3, max_value=10, value=4)
        
        # 2. PROSES PEMBAGIAN STRATA OTOMATIS (BINNING)
        try:
            # Gunakan qcut untuk membagi data berdasarkan persentil
            # duplicates='drop' berguna jika banyak angka yang sama persis (misal banyak nilai 0)
            df['Strata'], bin_edges = pd.qcut(df[value_col], q=n_bins, retbins=True, duplicates='drop', precision=0)
            
            # Mempercantik Label Strata agar lebih mudah dibaca user
            # Kita ubah label default pandas yang kaku menjadi format Rupiah
            new_labels = []
            # qcut mungkin mengurangi jumlah bin jika data tidak unik, kita cek jumlah kategori aslinya
            cat_codes = df['Strata'].cat.categories
            
            for cat in cat_codes:
                # Format label: "Rp 0 - Rp 1.000.000"
                lbl = f"{cat.left:,.0f} s.d {cat.right:,.0f}"
                new_labels.append(lbl)
            
            # Terapkan label baru
            df['Strata'] = df['Strata'].cat.rename_categories(new_labels)
            
            # Hitung Statistik per Strata (Count & StdDev)
            strata_stats = df.groupby('Strata', observed=True)[value_col].agg(['count', 'std', 'mean']).reset_index()
            
            # --- TAMPILAN TABEL EDITOR ---
            st.write("📊 Distribusi Populasi per Kuartil:")
            
            # Isi NaN pada std dengan 0 (jika populasi strata cuma 1, std tidak bisa dihitung)
            strata_stats['std'] = strata_stats['std'].fillna(0)
            
            strata_summary = []
            
            # Editor interaktif
            edited_strata = st.data_editor(
                strata_stats,
                column_config={
                    "Strata": "Range Nilai (Rupiah)",
                    "count": "Populasi (N)",
                    "mean": "Rata-rata",
                    "std": st.column_config.NumberColumn(
                        "Estimasi SD",
                        help="Standar Deviasi (Bisa diedit)",
                        required=True,
                        format="%.2f"
                    )
                },
                disabled=["Strata", "count", "mean"],
                hide_index=True,
                use_container_width=True
            )
            
            # Simpan hasil editan user ke list dictionary
            for index, row in edited_strata.iterrows():
                strata_summary.append({
                    'strata': row['Strata'],
                    'count': row['count'],
                    'std_dev': row['std']
                })
            
            # --- HITUNG JUMLAH SAMPEL (n) ---
            n_res, error_msg = calc.calculate_mpu_stratified(strata_summary, confidence, sst)
            
            # --- ALOKASI SAMPEL (SAMA SEPERTI KODE SEBELUMNYA) ---
            if n_res > 0:
                st.markdown("#### 📍 Alokasi Sampel Final")
                
                total_weight = sum([s['count'] * s['std_dev'] for s in strata_summary])
                allocation_dict = {} 
                final_allocations = []
                allocation_text = ""
                
                for s in strata_summary:
                    weight = s['count'] * s['std_dev']
                    
                    if total_weight > 0:
                        raw_allocation = (weight / total_weight) * n_res
                        n_teoritis = math.ceil(raw_allocation)
                    else:
                        n_teoritis = 0
                    
                    # Logika Clamping (Sensus jika n > Populasi)
                    if n_teoritis >= s['count']:
                        n_final_strata = s['count']
                        keterangan = "✅ **Sensus**"
                    else:
                        n_final_strata = n_teoritis
                        keterangan = ""
                    
                    final_allocations.append(n_final_strata)
                    allocation_dict[s['strata']] = int(n_final_strata)
                    
                    allocation_text += f"- **{s['strata']}** (Pop: {s['count']}): Ambil **{n_final_strata}** {keterangan}\n"
                
                st.markdown(allocation_text)
                
                # Update n_res total
                n_adjusted_total = sum(final_allocations)
                st.info(f"💡 Total Sampel: **{n_adjusted_total}** item.")
                n_res = n_adjusted_total
                
                # SIMPAN KE SESSION STATE
                st.session_state['allocation_dict'] = allocation_dict
                st.session_state['df_stratified'] = df

        except Exception as e:
            st.error(f"Gagal membagi kuartil otomatis. Data mungkin terlalu sedikit atau seragam. Error: {e}")
            n_res = 0

    elif metode_sampling == "Difference/Ratio Estimation":
        with col_in3:
            var = st.number_input("Estimasi Varians", value=1000000.0)
        n_res, error_msg = calc.calculate_difference_ratio(
            len(df), confidence, sst, var)

    # Tampilkan Hasil N
    if error_msg:
        st.error(error_msg)
        n_final = 0
    else:
        st.metric("Jumlah Sampel Disarankan (n)", f"{n_res} Item")
        n_final = st.number_input("Jumlah Sampel Final",
                                  min_value=1,
                                  max_value=len(df),
                                  value=int(n_res) if n_res > 0 else 30)

    # 4. TEKNIK EKSEKUSI
    st.markdown("---")
    st.subheader("Teknik Pemilihan Sampel")

    teknik = st.selectbox("Teknik Pemilihan", [
        "Acak Sederhana", "PPS (Wajib untuk MUS)", "Sistematis",
        "Sistematis Acak", "Stratifikasi (Top Value)",
        "Benford's Law (Anomali)"
    ])

    # --- PERBAIKAN: Gunakan Session State untuk tombol ---
    generate_btn = st.button("🚀 Generate Sampel")

    # Logika Hitung Sampel (Hanya jalan saat tombol diklik)
    if generate_btn:
        sampled_df = pd.DataFrame()

        # --- LOGIKA BARU: CEK APAKAH PAKAI STRATIFIED MPU ---
        if metode_sampling == "Stratified Mean Per Unit (MPU)" and 'allocation_dict' in st.session_state:
            st.info("Menggunakan pemilihan terdistribusi sesuai Strata...")
            
            # Ambil DF yang sudah ada label Strata-nya
            df_to_use = st.session_state.get('df_stratified', df)
            alloc_data = st.session_state['allocation_dict']

            # Panggil fungsi baru di selections.py
            sampled_df = sel.select_stratified_distributed(
                df_to_use,
                alloc_data,
                teknik, 
                value_col)

        else:
            # --- LOGIKA LAMA (UNTUK METODE NON-STRATIFIED) ---
            if teknik == "Acak Sederhana":
                sampled_df = sel.select_simple_random(df, n_final)
            elif teknik == "PPS (Wajib untuk MUS)":
                sampled_df = sel.select_pps(df, n_final, value_col)
            elif teknik == "Sistematis":
                sampled_df = sel.select_systematic(df, n_final)
            elif teknik == "Sistematis Acak":
                sampled_df = sel.select_random_systematic(df, n_final)
            elif teknik == "Stratifikasi (Top Value)":
                sampled_df = sel.select_stratified_top_value(df, n_final, value_col)
            elif teknik == "Benford's Law (Anomali)":
                sampled_df = sel.select_benford_anomaly(df, n_final, value_col)
        
        # SIMPAN HASIL KE SESSION STATE AGAR TIDAK HILANG SAAT REFRESH
        st.session_state['sampled_df'] = sampled_df
        # Reset state generate laporan agar tombol download docx hilang dulu jika generate ulang
        if 'report_docx_bytes' in st.session_state:
            del st.session_state['report_docx_bytes']

    # --- BAGIAN TAMPILAN (Dijalankan cek Session State, bukan tombol) ---
    if 'sampled_df' in st.session_state and not st.session_state['sampled_df'].empty:
        
        # Ambil data dari memory
        current_sampled_df = st.session_state['sampled_df']
        
        st.success(f"Terpilih {len(current_sampled_df)} sampel.")

        # Tampilkan rincian per strata (opsional)
        if 'Strata' in current_sampled_df.columns:
            st.write("Rincian Sampel per Strata:")
            st.write(current_sampled_df['Strata'].value_counts())

        st.dataframe(current_sampled_df)
        
        # === BUTTONS DOWNLOAD ===
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        with col_btn1:
            # Download Sampel Only
            buff_sampel = BytesIO()
            with pd.ExcelWriter(buff_sampel, engine='openpyxl') as writer:
                current_sampled_df.to_excel(writer, index=False, sheet_name='Sampel')
            buff_sampel.seek(0)
            st.download_button(
                label="📊 Download Sampel (.xlsx)",
                data=buff_sampel.getvalue(),
                file_name="sampel.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        with col_btn2:
            # Download Laporan Lengkap (.xlsx)
            buff_laporan = generate_laporan_xlsx(
                df,
                current_sampled_df,
                metode_sampling,
                teknik,
                confidence,
                sst,
                value_col,
                n_final
            )
            st.download_button(
                label="📋 Generate Laporan Lengkap (.xlsx)",
                data=buff_laporan.getvalue(),
                file_name=f"Laporan_Sampling_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        with col_btn3:
            # Generate & Download Laporan DOCX
            # Fungsi callback agar tidak refresh halaman total
            def on_generate_docx_inline():
                try:
                    # Ambil sampled_df dari session state
                    sdf = st.session_state['sampled_df']
                    buf = generate_laporan_docx(
                        df_original=df, # df diambil dari global scope (hasil upload)
                        sampled_df=sdf,
                        metode_sampling=metode_sampling,
                        teknik=teknik,
                        confidence=confidence,
                        sst=sst,
                        value_col=value_col,
                        n_final=n_final
                    )
                    st.session_state['report_docx_bytes'] = buf.getvalue()
                except Exception as exc:
                    st.error(f"❌ Gagal membuat .docx: {exc}")

            # Tombol pemicu generate docx
            st.button("📄 Generate Laporan (.docx)", on_click=on_generate_docx_inline, key="btn_docx_inline")
            
            # Tampilkan tombol download file-nya JIKA bytes sudah ada di session state
            if 'report_docx_bytes' in st.session_state and st.session_state['report_docx_bytes']:
                st.success("✅ Laporan .docx siap!")
                st.download_button(
                    label="⬇️ Download File (.docx)",
                    data=st.session_state['report_docx_bytes'],
                    file_name=f"Laporan_Sampling_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx_inline"
                )

    elif 'sampled_df' in st.session_state and st.session_state['sampled_df'].empty:
        st.warning("Tidak ada sampel yang terpilih. Cek parameter.")

else:
    st.info("Silakan upload data di sidebar.")


