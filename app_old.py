import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO

# --- IMPORT MODUL SENDIRI ---
import calculations as calc
import selections as sel
import pendapatan_analyzer as pend_analyzer
from helpers import generate_laporan_pendapatan_xlsx, generate_laporan_pendapatan_docx, generate_template_pendapatan
from belanja_dashboard import dashboard_belanja

# --- FUNGSI HELPER ---
def detect_encoding(uploaded_file):
    """
    Deteksi encoding file CSV.
    Coba: utf-8, latin-1, cp1252, iso-8859-1, windows-1252
    """
    try:
        uploaded_file.seek(0)
        raw_data = uploaded_file.read()
        
        # Gunakan chardet untuk deteksi encoding
        result = chardet.detect(raw_data)
        encoding = result.get('encoding', 'utf-8')
        
        uploaded_file.seek(0)
        return encoding if encoding else 'utf-8'
    except Exception as e:
        # Fallback ke encoding umum Windows
        uploaded_file.seek(0)
        return 'latin-1'

def detect_csv_delimiter(uploaded_file, sample_size=5000):
    """
    Deteksi delimiter CSV secara otomatis.
    Coba: ',', ';', '\t', '|'
    """
    encoding = detect_encoding(uploaded_file)
    
    try:
        # Baca sample dari file dengan encoding yang terdeteksi
        uploaded_file.seek(0)
        sample = uploaded_file.read(sample_size).decode(encoding, errors='ignore')
        
        # Coba deteksi dengan csv.Sniffer
        delimiter = csv.Sniffer().sniff(sample).delimiter
        uploaded_file.seek(0)
        return delimiter, encoding
    except Exception as e:
        # Jika gagal, coba delimiter umum
        delimiters = [';', ',', '\t', '|']
        
        for delim in delimiters:
            try:
                uploaded_file.seek(0)
                test_df = pd.read_csv(uploaded_file, sep=delim, nrows=5, encoding=encoding)
                if len(test_df.columns) > 1:
                    uploaded_file.seek(0)
                    return delim, encoding
            except:
                continue
        
        uploaded_file.seek(0)
        return ',', encoding  # Default ke comma


def convert_rupiah_to_numeric(df):
    """
    Konversi kolom Rupiah ke numerik dengan perlindungan ketat (Cek 5 Baris).
    Hanya konversi jika TIDAK ADA huruf sama sekali pada sampel data.
    """
    import re # Wajib import di sini
    
    print("--- Memulai Cek Konversi Data ---")

    for col in df.columns:
        try:
            # 1. Hanya proses kolom tipe Object (String)
            if df[col].dtype != 'object':
                continue

            # 2. Ambil 5 sampel data teratas yang tidak kosong (NaN)
            # Mengambil 5 baris lebih aman daripada cuma 1 baris
            valid_samples = df[col].dropna().head(5).astype(str).tolist()
            
            if not valid_samples:
                continue

            # 3. DETEKSI TEKS (HURUF)
            # Jika ada SATU SAJA sampel yang mengandung huruf a-z, skip kolom ini!
            is_text_column = False
            for s in valid_samples:
                # Bersihkan simbol umum, sisakan huruf dan angka
                clean_check = s.lower().replace('rp', '').replace('.', '').replace(',', '').replace('-', '').strip()
                
                # Cek regex huruf a-z
                if re.search('[a-z]', clean_check):
                    is_text_column = True
                    break # Stop, ini pasti kolom Nama/Keterangan
            
            if is_text_column:
                # print(f"ℹ️ Skip kolom '{col}': Terdeteksi sebagai Teks/Nama")
                continue

            # 4. Jika lolos cek huruf, baru coba konversi angka
            # Ambil sampel pertama untuk penentuan pola (Titik/Koma)
            sample = valid_samples[0]
            
            # Pola 1: Format Rupiah Indonesia (Ribuan Titik, Desimal Koma) -> 1.250.000,00
            if '.' in sample and ',' in sample:
                last_dot = sample.rfind('.')
                last_comma = sample.rfind(',')
                
                if last_dot < last_comma:
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace('Rp', '', regex=False)
                               .str.replace(' ', '', regex=False)
                               .str.replace('.', '', regex=False) # Hapus ribuan
                               .str.replace(',', '.', regex=False) # Koma jadi titik
                               .apply(pd.to_numeric, errors='coerce'))
                    print(f"✅ Konversi '{col}': Format Rupiah (Indo)")

            # Pola 2: Format Desimal Koma Saja (1250000,00)
            elif ',' in sample and '.' not in sample:
                # Pastikan yang tersisa hanya angka
                clean_sample = sample.replace(',', '').strip()
                if clean_sample.isdigit():
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace(',', '.', regex=False)
                               .apply(pd.to_numeric, errors='coerce'))
                    print(f"✅ Konversi '{col}': Desimal Koma")

        except Exception as e:
            # Jika error, biarkan kolom apa adanya
            pass
    
    print("--- Selesai Cek Konversi ---")
    return df


# --- FUNGSI UNTUK TEMPLATE PENDAPATAN ---
def generate_template_pendapatan():
    """
    Generate template Excel untuk analisis Pendapatan WP.
    Kolom: Nomor, Nama WP, NPWPD, JANUARI-DESEMBER
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Template Pendapatan"
    
    # Header
    headers = ['NOMOR', 'NAMA WP', 'NPWPD'] + [
        'JANUARI', 'FEBRUARI', 'MARET', 'APRIL', 'MEI', 'JUNI',
        'JULI', 'AGUSTUS', 'SEPTEMBER', 'OKTOBER', 'NOVEMBER', 'DESEMBER'
    ]
    
    ws.append(headers)
    
    # Format header
    fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    font_header = Font(bold=True, color="FFFFFF", size=11)
    alignment_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    alignment_right = Alignment(horizontal="right", vertical="center")
    
    for cell in ws[1]:
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = alignment_center
    
    # Tambah contoh data (opsional)
    ws.append([1, 'Contoh WP 1', '12.345.678.901-123.456', 1000000, 1000000, 1000000, 1000000, 
               1000000, 1000000, 1000000, 1000000, 1000000, 1000000, 1000000, 1000000])
    
    # Set lebar kolom
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 20
    
    for col in ['D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O']:
        ws.column_dimensions[col].width = 14
        # Format nilai sebagai currency
        for row in ws.iter_rows(min_col=ws[col+'1'].column, max_col=ws[col+'1'].column, min_row=2):
            for cell in row:
                cell.alignment = alignment_right
    
    buff = BytesIO()
    wb.save(buff)
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK LAPORAN PENDAPATAN (EXCEL) ---
def generate_laporan_pendapatan_xlsx(df_original, anomali_list, bulan_cols):
    """
    Generate laporan analisis anomali Pendapatan dalam format Excel.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, numbers
    
    buff = BytesIO()
    
    # Hitung statistik
    statistik = pend_analyzer.hitung_statistik_pendapatan(df_original, bulan_cols)
    
    with pd.ExcelWriter(buff, engine='openpyxl') as writer:
        # === SHEET 1: RINGKASAN ===
        ringkasan_data = {
            'Keterangan': [
                'Total WP dalam Populasi',
                'WP dengan Anomali Terdeteksi',
                'Persentase WP Anomali (%)',
                'Total Pendapatan (Rp)',
                'Rata-rata Pendapatan (Rp)',
                'Median Pendapatan (Rp)',
                'Standar Deviasi (Rp)',
                'Pendapatan Minimum (Rp)',
                'Pendapatan Maksimum (Rp)',
                'Tanggal Generate Laporan'
            ],
            'Nilai': [
                statistik['total_wp'],
                len(anomali_list),
                f"{(len(anomali_list) / statistik['total_wp'] * 100):.2f}%" if statistik['total_wp'] > 0 else "0%",
                f"Rp {statistik['total_pendapatan']:,.2f}",
                f"Rp {statistik['rata_rata']:,.2f}",
                f"Rp {statistik['median']:,.2f}",
                f"Rp {statistik['std_dev']:,.2f}",
                f"Rp {statistik['min']:,.2f}",
                f"Rp {statistik['max']:,.2f}",
                pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S")
            ]
        }
        df_ringkasan = pd.DataFrame(ringkasan_data)
        df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
        
        # === SHEET 2: DAFTAR ANOMALI ===
        if anomali_list:
            df_anomali = pd.DataFrame(anomali_list)
            # Rename kolom untuk tampilan yang lebih baik
            df_anomali = df_anomali.rename(columns={
                'nomor': 'No',
                'nama_wp': 'Nama WP',
                'npwpd': 'NPWPD',
                'jenis_anomali': 'Jenis Anomali',
                'bulan_terisi': 'Bulan Terisi',
                'rata_rata': 'Rata-rata (Rp)',
                'min': 'Min (Rp)',
                'max': 'Max (Rp)',
                'std_dev': 'Std Dev (Rp)'
            })
            df_anomali.to_excel(writer, sheet_name='Daftar Anomali', index=False)
        
        # === SHEET 3: BREAKDOWN JENIS ANOMALI ===
        if anomali_list:
            breakdown_data = []
            jenis_anomali_set = set()
            for item in anomali_list:
                for jenis in item['jenis_anomali'].split(' | '):
                    jenis_anomali_set.add(jenis)
            
            for jenis in jenis_anomali_set:
                count = sum(1 for item in anomali_list if jenis in item['jenis_anomali'])
                breakdown_data.append({
                    'Jenis Anomali': jenis,
                    'Jumlah WP': count,
                    'Persentase (%)': f"{(count / len(anomali_list) * 100):.2f}%"
                })
            
            df_breakdown = pd.DataFrame(breakdown_data)
            df_breakdown.to_excel(writer, sheet_name='Breakdown Anomali', index=False)
        
        # === FORMATTING ===
        workbook = writer.book
        
        # Format Ringkasan
        ws_ringkasan = workbook['Ringkasan']
        fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        font_header = Font(bold=True, color="FFFFFF")
        
        for cell in ws_ringkasan[1]:
            cell.fill = fill_header
            cell.font = font_header
        
        for column in ws_ringkasan.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws_ringkasan.column_dimensions[column_letter].width = (max_length + 2)
        
        # Format Daftar Anomali
        if 'Daftar Anomali' in workbook.sheetnames:
            ws_anomali = workbook['Daftar Anomali']
            for cell in ws_anomali[1]:
                cell.fill = fill_header
                cell.font = font_header
    
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK LAPORAN PENDAPATAN (DOCX) ---
def generate_laporan_pendapatan_docx(df_original, anomali_list, bulan_cols):
    """
    Generate laporan analisis anomali Pendapatan dalam format Word.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # Judul
    h = doc.add_heading('Laporan Analisis Anomali Pendapatan Wajib Pajak', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tanggal Generate: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S"))
    
    # Statistik Umum
    statistik = pend_analyzer.hitung_statistik_pendapatan(df_original, bulan_cols)
    
    doc.add_heading('Ringkasan Hasil Analisis', level=2)
    doc.add_paragraph(f"Total WP dalam Populasi: {statistik['total_wp']} WP")
    doc.add_paragraph(f"WP dengan Anomali Terdeteksi: {len(anomali_list)} WP")
    doc.add_paragraph(f"Persentase WP Anomali: {(len(anomali_list) / statistik['total_wp'] * 100):.2f}%")
    
    doc.add_heading('Statistik Pendapatan', level=2)
    doc.add_paragraph(f"Total Pendapatan: Rp {statistik['total_pendapatan']:,.2f}")
    doc.add_paragraph(f"Rata-rata Pendapatan: Rp {statistik['rata_rata']:,.2f}")
    doc.add_paragraph(f"Median Pendapatan: Rp {statistik['median']:,.2f}")
    doc.add_paragraph(f"Standar Deviasi: Rp {statistik['std_dev']:,.2f}")
    doc.add_paragraph(f"Pendapatan Minimum: Rp {statistik['min']:,.2f}")
    doc.add_paragraph(f"Pendapatan Maksimum: Rp {statistik['max']:,.2f}")
    
    # Breakdown Anomali
    if anomali_list:
        doc.add_heading('Breakdown Jenis Anomali', level=2)
        jenis_anomali_count = {}
        for item in anomali_list:
            for jenis in item['jenis_anomali'].split(' | '):
                jenis_anomali_count[jenis] = jenis_anomali_count.get(jenis, 0) + 1
        
        for jenis, count in jenis_anomali_count.items():
            pct = (count / len(anomali_list) * 100)
            doc.add_paragraph(f"{jenis}: {count} WP ({pct:.2f}%)")
    
    # Daftar WP dengan Anomali
    doc.add_heading('Daftar WP dengan Anomali', level=2)
    
    if anomali_list:
        # Buat tabel
        tbl = doc.add_table(rows=1, cols=5)
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Nama WP'
        hdr_cells[2].text = 'NPWPD'
        hdr_cells[3].text = 'Jenis Anomali'
        hdr_cells[4].text = 'Rata-rata Pendapatan'
        
        for item in anomali_list[:100]:  # Limit 100 baris untuk dokumen
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(item['nomor'])
            row_cells[1].text = str(item['nama_wp'])
            row_cells[2].text = str(item['npwpd'])
            row_cells[3].text = str(item['jenis_anomali'])
            row_cells[4].text = f"Rp {item['rata_rata']:,.2f}"
        
        if len(anomali_list) > 100:
            doc.add_paragraph(f"... dan {len(anomali_list) - 100} WP lainnya (lihat laporan Excel untuk lengkapnya)")
    else:
        doc.add_paragraph("Tidak ditemukan WP dengan anomali berdasarkan kriteria analisis.")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

    

#KODE untuk buat Laporan dalam bentuk excel
def generate_laporan_xlsx(df_original, sampled_df, metode_sampling, teknik, 
                          confidence, sst, value_col, n_final):
    """
    Generate laporan hasil sampling dalam format Excel (.xlsx)
    dengan multiple sheets: Ringkasan, Metode, Detail Sampel
    """
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    buff = BytesIO()
    
    with pd.ExcelWriter(buff, engine='openpyxl') as writer:
        # === SHEET 1: RINGKASAN ===
        # Hitung total nilai sampel dan persentasenya
        total_nilai_buku = df_original[value_col].sum()
        total_nilai_sampel = sampled_df[value_col].sum() if not sampled_df.empty else 0
        pct_nilai_sampel = (total_nilai_sampel / total_nilai_buku * 100) if total_nilai_buku > 0 else 0
        
        ringkasan_data = {
            'Keterangan': [
                'Total Populasi (N)',
                'Total Nilai Buku (Rp)',
                'Total Nilai Sampel (Rp)',
                'Persentase Nilai Sampel (%)',
                'Jumlah Sampel (n)',
                'Persentase Sampel (%)',
                'Metode Sampling',
                'Teknik Pemilihan',
                'Confidence Level (%)',
                'Salah Saji Tertoleransi (Rp)',
                'Tanggal Generate Laporan'
            ],
            'Nilai': [
                len(df_original),
                f"Rp {total_nilai_buku:,.2f}",
                f"Rp {total_nilai_sampel:,.2f}",
                f"{pct_nilai_sampel:.2f}%",
                n_final,
                f"{(n_final / len(df_original) * 100):.2f}%",
                metode_sampling,
                teknik,
                confidence,
                f"Rp {sst:,.2f}",
                pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S")
            ]
        }
        df_ringkasan = pd.DataFrame(ringkasan_data)
        df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
        
        # === SHEET 2: DETAIL SAMPEL ===
        sampled_df.to_excel(writer, sheet_name='Detail Sampel', index=False)
        
        # === SHEET 3: STATISTIK POPULASI ===
        statistik_data = {
            'Metrik': [
                'Mean',
                'Median',
                'Std Dev',
                'Min',
                'Max',
                'Q1 (25%)',
                'Q3 (75%)'
            ],
            'Populasi': [
                f"Rp {df_original[value_col].mean():,.2f}",
                f"Rp {df_original[value_col].median():,.2f}",
                f"Rp {df_original[value_col].std():,.2f}",
                f"Rp {df_original[value_col].min():,.2f}",
                f"Rp {df_original[value_col].max():,.2f}",
                f"Rp {df_original[value_col].quantile(0.25):,.2f}",
                f"Rp {df_original[value_col].quantile(0.75):,.2f}"
            ],
            'Sampel': [
                f"Rp {sampled_df[value_col].mean():,.2f}",
                f"Rp {sampled_df[value_col].median():,.2f}",
                f"Rp {sampled_df[value_col].std():,.2f}",
                f"Rp {sampled_df[value_col].min():,.2f}",
                f"Rp {sampled_df[value_col].max():,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.25):,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.75):,.2f}"
            ]
        }
        df_statistik = pd.DataFrame(statistik_data)
        df_statistik.to_excel(writer, sheet_name='Statistik', index=False)
        
        # === FORMATTING ===
        workbook = writer.book
        
        # Format Ringkasan
        ws_ringkasan = workbook['Ringkasan']
        fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        font_header = Font(bold=True, color="FFFFFF")
        
        for cell in ws_ringkasan[1]:
            cell.fill = fill_header
            cell.font = font_header
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Auto width
        for column in ws_ringkasan.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws_ringkasan.column_dimensions[column_letter].width = adjusted_width
        
        # Format Detail Sampel
        ws_detail = workbook['Detail Sampel']
        for cell in ws_detail[1]:
            cell.fill = fill_header
            cell.font = font_header
        
        # Format Statistik
        ws_statistik = workbook['Statistik']
        for cell in ws_statistik[1]:
            cell.fill = fill_header
            cell.font = font_header
    
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK GENERATE LAPORAN DOCX ---
def generate_laporan_docx(df_original, sampled_df, metode_sampling, teknik,
                          confidence, sst, value_col, n_final, max_rows=300):
    """
    Buat laporan .docx:
    - Ringkasan metadata
    - Statistik populasi vs sampel
    - Potongan Detail Sampel (maks max_rows)
    Mengembalikan BytesIO yang siap di-download.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Judul
    h = doc.add_heading('Laporan Sampling Pemeriksaan Keuangan', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata / Ringkasan singkat
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tanggal: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S"))
    doc.add_paragraph(f"Metode Sampling: {metode_sampling}")
    doc.add_paragraph(f"Teknik Pemilihan: {teknik}")
    doc.add_paragraph(f"Confidence Level: {confidence}%")
    doc.add_paragraph(f"Salah Saji Tertoleransi (SST): Rp {sst:,.2f}")
    doc.add_paragraph(f"Jumlah Populasi (N): {len(df_original)}")
    if value_col:
        try:
            total_nilai_buku = df_original[value_col].sum()
            doc.add_paragraph(f"Total Nilai Buku: Rp {total_nilai_buku:,.2f}")
            total_nilai_sampel = sampled_df[value_col].sum() if not sampled_df.empty else 0
            pct_nilai = (total_nilai_sampel / total_nilai_buku * 100) if total_nilai_buku > 0 else 0
            doc.add_paragraph(f"Total Nilai Sampel: Rp {total_nilai_sampel:,.2f}")
            doc.add_paragraph(f"Persentase Nilai Sampel: {pct_nilai:.2f}%")
        except Exception:
            doc.add_paragraph("Total Nilai Buku: -")
    doc.add_paragraph(f"Jumlah Sampel (n): {n_final}")
    doc.add_paragraph()

    # Statistik: buat tabel metric vs populasi vs sampel
    metrics = ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Q1 (25%)', 'Q3 (75%)']
    pop_values = []
    samp_values = []
    if value_col and value_col in df_original.columns and value_col in sampled_df.columns:
        pop = df_original[value_col].dropna().astype(float)
        samp = sampled_df[value_col].dropna().astype(float)
        pop_values = [
            f"Rp {pop.mean():,.2f}",
            f"Rp {pop.median():,.2f}",
            f"Rp {pop.std():,.2f}",
            f"Rp {pop.min():,.2f}",
            f"Rp {pop.max():,.2f}",
            f"Rp {pop.quantile(0.25):,.2f}",
            f"Rp {pop.quantile(0.75):,.2f}"
        ]
        samp_values = [
            f"Rp {samp.mean():,.2f}" if not samp.empty else "-",
            f"Rp {samp.median():,.2f}" if not samp.empty else "-",
            f"Rp {samp.std():,.2f}" if not samp.empty else "-",
            f"Rp {samp.min():,.2f}" if not samp.empty else "-",
            f"Rp {samp.max():,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.25):,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.75):,.2f}" if not samp.empty else "-"
        ]
    else:
        pop_values = ['-'] * len(metrics)
        samp_values = ['-'] * len(metrics)

    doc.add_heading('Statistik Populasi vs Sampel', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = 'Metrik'
    hdr_cells[1].text = 'Populasi'
    hdr_cells[2].text = 'Sampel'
    for i, m in enumerate(metrics):
        row_cells = tbl.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = pop_values[i]
        row_cells[2].text = samp_values[i]

    doc.add_paragraph()

    # Detail Sampel (ringkasan)
    doc.add_heading('Detail Sampel', level=2)
    if sampled_df.empty:
        doc.add_paragraph("Tidak ada sampel terpilih.")
    else:
        # Beri keterangan agar pembaca membuka lampiran Excel untuk daftar lengkap
        doc.add_paragraph(
            "Daftar sampel selengkapnya dilampirkan pada file Excel laporan. "
            "Untuk melihat data lengkap (semua kolom dan baris), silakan unduh laporan .xlsx."
        )


    # Simpan ke BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Callback yang dipanggil saat tombol Generate .docx diklik
def on_generate_docx():
    try:
        sampled_df = st.session_state.get('sampled_df', pd.DataFrame())
        if sampled_df is None or sampled_df.empty:
            st.warning("Belum ada sampel. Klik 'Generate Sampel' terlebih dahulu.")
            return

        # panggil fungsi pembuat docx yang Anda miliki; pastikan mengembalikan BytesIO
        buf = generate_laporan_docx(
            df_original=df,             # pastikan `df` adalah dataframe populasi yang ada di scope
            sampled_df=sampled_df,
            metode_sampling=st.session_state.get('sampled_params', {}).get('metode_sampling', metode_sampling),
            teknik=st.session_state.get('sampled_params', {}).get('teknik', teknik),
            confidence=confidence,
            sst=sst,
            value_col=value_col,
            n_final=n_final
        )
        # simpan bytes ke session_state agar tidak dihitung ulang
        st.session_state['report_docx_bytes'] = buf.getvalue()
        st.success("Laporan .docx berhasil dibuat. Tombol download muncul di bawah.")
    except Exception as exc:
        st.error(f"Gagal membuat laporan .docx: {exc}")
        # untuk debugging, tunjukkan traceback singkat (opsional)
        import traceback
        st.text(traceback.format_exc())


st.set_page_config(page_title="Dashboard Sampling Audit", layout="wide")
st.title("🕵️ Dashboard Uji Petik Pemeriksaan Keuangan")
st.markdown("---")

# === PILIHAN JENIS SAMPLING ===
jenis_sampling = st.selectbox(
    "📋 Pilih Jenis Sampling:",
    ["Belanja/Lainnya", "Pendapatan"],
    help="Pilih jenis sampling yang akan dianalisis"
)

st.markdown("---")

# === CONDITIONAL LOGIC BERDASARKAN JENIS SAMPLING ===
if jenis_sampling == "Pendapatan":
    # ===== DASHBOARD PENDAPATAN =====
    st.subheader("📊 Analisis Anomali Pendapatan Wajib Pajak")
    st.write("Unggah data pelaporan untuk mengidentifikasi WP dengan pelaporan identik atau variasi rendah (<=20%).")
    
    # Download Template
    st.markdown("---")
    st.subheader("📥 Download Template Kertas Kerja")
    st.write("Silakan download template di bawah sebagai format standard untuk upload data:")
    
    template_buff = generate_template_pendapatan()
    st.download_button(
        label="📄 Download Template Pendapatan",
        data=template_buff.getvalue(),
        file_name="Template_Pendapatan.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    
    st.markdown("---")
    st.subheader("📤 Upload Data Pendapatan")
    
    # Upload File
    uploaded_file_pend = st.file_uploader(
        "Pilih file (Excel, CSV, atau Parquet)",
        type=['xlsx', 'csv', 'parquet'],
        key='pend_uploader'
    )
    
    if uploaded_file_pend:
        # Load data berdasarkan ekstensi
        try:
            if uploaded_file_pend.name.endswith('.csv'):
                df_pend = pd.read_csv(uploaded_file_pend)
            elif uploaded_file_pend.name.endswith('.parquet'):
                df_pend = pd.read_parquet(uploaded_file_pend)
            else:
                df_pend = pd.read_excel(uploaded_file_pend)
            
            st.success(f"✅ Data berhasil dimuat: {len(df_pend)} baris")
            
            # Preview Data
            with st.expander("👁️ Preview Data", expanded=False):
                st.dataframe(df_pend.head())
            
            # Tentukan kolom bulan yang tersedia
            bulan_list = ['JANUARI', 'FEBRUARI', 'MARET', 'APRIL', 'MEI', 'JUNI',
                          'JULI', 'AGUSTUS', 'SEPTEMBER', 'OKTOBER', 'NOVEMBER', 'DESEMBER']
            available_bulan = [b for b in bulan_list if b in df_pend.columns]
            
            if available_bulan:
                st.info(f"📅 Kolom bulan terdeteksi: {', '.join(available_bulan)}")
                
                # Jalankan analisis
                if st.button("🔍 Lakukan Analisis Anomali", key="btn_analisis_pend"):
                    with st.spinner("Menganalisis anomali pendapatan..."):
                        anomali_results = pend_analyzer.detect_anomali_pendapatan(df_pend, available_bulan)
                    
                    st.session_state['anomali_results'] = anomali_results
                    st.session_state['df_pendapatan'] = df_pend
                    st.session_state['bulan_cols_pend'] = available_bulan
                
                # Tampilkan hasil jika sudah ada
                if 'anomali_results' in st.session_state:
                    anomali_results = st.session_state['anomali_results']
                    df_pend = st.session_state['df_pendapatan']
                    available_bulan = st.session_state.get('bulan_cols_pend', available_bulan)
                    
                    st.markdown("---")
                    st.subheader("📈 Hasil Analisis")
                    
                    # Hitung statistik untuk display
                    stat_pend = pend_analyzer.hitung_statistik_pendapatan(df_pend, available_bulan)
                    
                    # Metrics
                    col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                    with col_m1:
                        st.metric("Total WP", stat_pend['total_wp'])
                    with col_m2:
                        st.metric("WP dengan Anomali", len(anomali_results))
                    with col_m3:
                        pct_anomali = (len(anomali_results) / stat_pend['total_wp'] * 100) if stat_pend['total_wp'] > 0 else 0
                        st.metric("% Anomali", f"{pct_anomali:.2f}%")
                    with col_m4:
                        st.metric("Total Pendapatan", f"Rp {stat_pend['total_pendapatan']:,.0f}")
                    
                    # Breakdown Anomali
                    if anomali_results:
                        st.subheader("📊 Breakdown Jenis Anomali")
                        jenis_anomali_count = {}
                        for item in anomali_results:
                            for jenis in item['jenis_anomali'].split(' | '):
                                jenis_anomali_count[jenis] = jenis_anomali_count.get(jenis, 0) + 1
                        
                        col_a1, col_a2 = st.columns(2)
                        with col_a1:
                            for jenis, count in jenis_anomali_count.items():
                                pct = (count / len(anomali_results) * 100) if anomali_results else 0
                                st.write(f"• **{jenis}**: {count} WP ({pct:.2f}%)")
                    
                    # Daftar Anomali
                    st.subheader("📋 Daftar WP dengan Anomali")
                    if anomali_results:
                        df_anomali_display = pd.DataFrame(anomali_results)[['nomor', 'nama_wp', 'npwpd', 'jenis_anomali', 'rata_rata']]
                        df_anomali_display = df_anomali_display.rename(columns={
                            'nomor': 'No',
                            'nama_wp': 'Nama WP',
                            'npwpd': 'NPWPD',
                            'jenis_anomali': 'Jenis Anomali',
                            'rata_rata': 'Rata-rata Pendapatan'
                        })
                        st.dataframe(df_anomali_display, use_container_width=True)
                        
                        # Download Laporan
                        st.markdown("---")
                        st.subheader("📥 Download Laporan")
                        
                        col_dl1, col_dl2 = st.columns(2)
                        
                        with col_dl1:
                            laporan_xlsx_pend = generate_laporan_pendapatan_xlsx(
                                df_pend, anomali_results, available_bulan
                            )
                            st.download_button(
                                label="📊 Download Laporan (.xlsx)",
                                data=laporan_xlsx_pend.getvalue(),
                                file_name=f"Laporan_Anomali_Pendapatan_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        
                        with col_dl2:
                            laporan_docx_pend = generate_laporan_pendapatan_docx(
                                df_pend, anomali_results, available_bulan
                            )
                            st.download_button(
                                label="📄 Download Laporan (.docx)",
                                data=laporan_docx_pend.getvalue(),
                                file_name=f"Laporan_Anomali_Pendapatan_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    else:
                        st.success("✅ Tidak ditemukan anomali berdasarkan kriteria analisis.")
            else:
                st.error("❌ Kolom bulan (JANUARI - DESEMBER) tidak ditemukan dalam file. Silakan gunakan template yang tersedia.")
        
        except Exception as e:
            st.error(f"❌ Gagal membaca file: {e}")

else:
    # ===== DASHBOARD BELANJA (ORIGINAL) =====
    # === PANDUAN METODE UJI PETIK ===
    with st.expander("📚 Panduan Metode Uji Petik", expanded=False):
    st.markdown("""
    ### Panduan Penentuan Metode Uji Petik
    
    Dalam menentukan metode uji petik, pemeriksa pada umumnya mempertimbangkan dua aspek utama, yaitu **nilai transaksi** dan **jumlah transaksi**. 
    
    - **Nilai transaksi** merupakan nilai uang yang diakui dan dicatat dalam laporan keuangan sebagai akibat dari suatu transaksi
    - **Jumlah transaksi** adalah banyaknya kejadian transaksi yang dicatat dalam suatu periode tertentu
    
    Berdasarkan kedua pertimbangan tersebut, pemeriksa dapat menentukan metode uji petik statistik yang sesuai dengan karakteristik populasi yang diperiksa:
    
    ---
    
    #### 1️⃣ **Monetary Unit Sampling (MUS)**
    
    Metode ini dapat digunakan apabila pemeriksa melakukan uji petik yang **berfokus pada nilai transaksi**. 
    
    Dalam metode ini, transaksi dengan nilai yang lebih besar memiliki peluang lebih tinggi untuk terpilih sebagai sampel.  
    
    ---
    
    #### 2️⃣ **Unstratified Mean Per Unit (MPU)**
    
    Metode ini dapat digunakan apabila pemeriksa melakukan uji petik yang **berfokus pada jumlah transaksi**, tanpa mempertimbangkan perbedaan nilai transaksi. 
    
    Metode ini sesuai diterapkan pada populasi dengan karakteristik nilai transaksi yang **relatif homogen** dan tidak memiliki rentang nilai yang terlalu besar.
    
    ---
    
    #### 3️⃣ **Stratified Mean Per Unit (MPU)**
    
    Metode ini dapat digunakan apabila pemeriksa **mempertimbangkan jumlah transaksi sekaligus nilai transaksi** dalam penentuan sampel. 
    
    Metode ini sesuai untuk populasi dengan karakteristik nilai transaksi yang **heterogen** serta memiliki **rentang nilai yang besar**, sehingga diperlukan pengelompokan (stratifikasi) untuk meningkatkan efektivitas dan representativitas sampel.
    
    """)

# 1. UPLOAD
st.sidebar.header("1. Upload Data")
uploaded_file = st.sidebar.file_uploader("Upload Tabel Data (Excel/CSV/Parquet)",
                                         type=['xlsx', 'csv', 'parquet'])

if uploaded_file is not None:
    # Load Data
    if uploaded_file.name.endswith('.csv'):
        try:
            # Deteksi delimiter dan encoding
            delimiter, encoding = detect_csv_delimiter(uploaded_file)
            df = pd.read_csv(uploaded_file, sep=delimiter, encoding=encoding)
            
            # 🔥 KONVERSI FORMAT RUPIAH KE NUMERIK
            df = convert_rupiah_to_numeric(df)
            
            st.info(f"📌 Delimiter: **'{delimiter}'** | Encoding: **{encoding}**")
        except Exception as e:
            st.error(f"❌ Gagal membaca file CSV: {e}")
            st.stop()
    
    elif uploaded_file.name.endswith('.parquet'):
        try:
            df = pd.read_parquet(uploaded_file)
            st.info("✅ File Parquet berhasil dibaca")
        except Exception as e:
            st.error(f"❌ Gagal membaca file Parquet: {e}")
            st.stop()
    
    else:  # Excel (.xlsx)
        try:
            df = pd.read_excel(uploaded_file)
            st.info("✅ File Excel berhasil dibaca")
        except Exception as e:
            st.error(f"❌ Gagal membaca file Excel: {e}")
            st.stop()

    st.success(f"Data dimuat: {len(df)} baris.")
    
    # untuk menampilkan preview data
    # st.write("📋 Preview Data:")
    # st.dataframe(df.head())
    # st.write("Tipe Data Kolom:")
    #st.write(df.dtypes)

    # Konversi kolom dengan format Rupiah ke numerik (pakai fungsi yang lebih ketat)
    # --- GANTI: gunakan fungsi `convert_rupiah_to_numeric` yang memeriksa beberapa sampel
    # agar kolom teks seperti nama, keterangan, pegawai tidak keliru dikonversi menjadi NaN.
    try:
        df = convert_rupiah_to_numeric(df)
    except Exception:
        # Jika gagal, biarkan dataframe apa adanya dan log untuk debugging
        print("⚠️ Gagal menjalankan convert_rupiah_to_numeric; melewatkan konversi.")
        pass
    
    # 2. PEMETAAN
    col1, col2 = st.columns(2)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    all_cols = df.columns.tolist()

    with col1:
        id_col = st.selectbox("Kolom ID Sampel", all_cols)
    with col2:
        value_col = st.selectbox("Kolom Nilai Rupiah", numeric_cols)

    total_nilai_buku = df[value_col].sum()
    st.info(f"💰 Total Nilai Buku: Rp {total_nilai_buku:,.2f}")

    # 3. METODE HITUNG JUMLAH SAMPEL
    st.markdown("---")
    st.subheader("Metode Penentuan Jumlah Sampel")

    metode_sampling = st.selectbox("Pilih Metode", [
        "Monetary Unit Sampling (MUS)", "Unstratified Mean Per Unit (MPU)",
        "Stratified Mean Per Unit (MPU)"
    ])

    # Parameter Input UI
    col_in1, col_in2, col_in3, col_in4 = st.columns(4)
    with col_in1:
        confidence = st.selectbox("Confidence Level", [90, 95, 99], index=1)
    with col_in2:
        sst = st.number_input("Salah Saji Tertoleransi (SST)",
                              value=float(total_nilai_buku) * 0.05)

    n_res = 0
    error_msg = None

    # --- PANGGIL FUNGSI DARI MODUL CALCULATIONS ---
    if metode_sampling == "Monetary Unit Sampling (MUS)":
        with col_in3:
            dss = st.number_input("Dugaan Salah Saji (DSS)", value=0.0)
        with col_in4:
            expansion = st.selectbox("Expansion Factor", [1, 5, 10, 15, 20, 25, 30, 37], index=1)
        n_res, error_msg = calc.calculate_mus(total_nilai_buku, confidence, sst, dss, expansion)

    elif metode_sampling == "Unstratified Mean Per Unit (MPU)":
        with col_in3:
            # --- UPDATE: HITUNG OTOMATIS SD DARI DATA ---
            # Sesuai Juknis: Jika tidak ada data tahun lalu, gunakan data saat ini (piloting)
            # Kita hitung SD dari kolom nilai rupiah di dataframe
            current_std_dev = df[value_col].std()
            
            st.info(f"💡 SD Populasi saat ini: {current_std_dev:,.2f}")
            
            sd = st.number_input(
                "Estimasi Standar Deviasi (SD)", 
                value=float(current_std_dev), # Default value langsung dari data
                min_value=0.0,
                format="%.2f",
                help="Default diisi dengan SD Populasi saat ini (sesuai saran Piloting Sample). Bisa diubah jika menggunakan data tahun lalu."
            )

        # Panggil fungsi perhitungan (Rumusnya sudah benar)
        n_res, error_msg = calc.calculate_mpu_unstratified(
            len(df), confidence, sst, sd
        )

    elif metode_sampling == "Stratified Mean Per Unit (MPU)":
        st.info("ℹ️ Stratifikasi otomatis menggunakan metode Kuantil (Membagi populasi sama rata).")
        
        # 1. KONSEP PILIHAN KUARTIL
        # Default audit biasanya 4 (Kuartil), tapi kita beri fleksibilitas
        n_bins = st.slider("Bagi populasi menjadi berapa bagian? (Default: 4 - Strata)", min_value=3, max_value=10, value=4)
        
        # 2. PROSES PEMBAGIAN STRATA OTOMATIS (BINNING)
        try:
            # Gunakan qcut untuk membagi data berdasarkan persentil
            # duplicates='drop' berguna jika banyak angka yang sama persis (misal banyak nilai 0)
            df['Strata'], bin_edges = pd.qcut(df[value_col], q=n_bins, retbins=True, duplicates='drop', precision=0)
            
            # Mempercantik Label Strata agar lebih mudah dibaca user
            # Kita ubah label default pandas yang kaku menjadi format Rupiah
            new_labels = []
            # qcut mungkin mengurangi jumlah bin jika data tidak unik, kita cek jumlah kategori aslinya
            cat_codes = df['Strata'].cat.categories
            
            for cat in cat_codes:
                # Format label: "Rp 0 - Rp 1.000.000"
                lbl = f"{cat.left:,.0f} s.d {cat.right:,.0f}"
                new_labels.append(lbl)
            
            # Terapkan label baru
            df['Strata'] = df['Strata'].cat.rename_categories(new_labels)
            
            # Hitung Statistik per Strata (Count & StdDev)
            strata_stats = df.groupby('Strata', observed=True)[value_col].agg(['count', 'std', 'mean']).reset_index()
            
            # --- TAMPILAN TABEL EDITOR ---
            st.write("📊 Distribusi Populasi per Strata:")
            
            # Isi NaN pada std dengan 0 (jika populasi strata cuma 1, std tidak bisa dihitung)
            strata_stats['std'] = strata_stats['std'].fillna(0)
            
            strata_summary = []
            
            # Editor interaktif
            edited_strata = st.data_editor(
                strata_stats,
                column_config={
                    "Strata": "Range Nilai (Rupiah)",
                    "count": "Populasi (N)",
                    "mean": "Rata-rata",
                    "std": st.column_config.NumberColumn(
                        "Estimasi SD",
                        help="Standar Deviasi (Bisa diedit)",
                        required=True,
                        format="%.2f"
                    )
                },
                disabled=["Strata", "count", "mean"],
                hide_index=True,
                use_container_width=True
            )
            
            # Simpan hasil editan user ke list dictionary
            for index, row in edited_strata.iterrows():
                strata_summary.append({
                    'strata': row['Strata'],
                    'count': row['count'],
                    'std_dev': row['std']
                })
            
            # --- HITUNG JUMLAH SAMPEL (n) ---
            n_res, error_msg = calc.calculate_mpu_stratified(strata_summary, confidence, sst)
            
            # --- ALOKASI SAMPEL (SAMA SEPERTI KODE SEBELUMNYA) ---
            if n_res > 0:
                st.markdown("#### 📍 Alokasi Sampel")
                
                total_weight = sum([s['count'] * s['std_dev'] for s in strata_summary])
                allocation_dict_default = {} 
                allocation_list_default = []
                
                for s in strata_summary:
                    weight = s['count'] * s['std_dev']
                    
                    if total_weight > 0:
                        raw_allocation = (weight / total_weight) * n_res
                        n_teoritis = math.ceil(raw_allocation)
                    else:
                        n_teoritis = 0
                    
                    # Logika Clamping (Sensus jika n > Populasi)
                    if n_teoritis >= s['count']:
                        n_final_strata = s['count']
                    else:
                        n_final_strata = n_teoritis
                    
                    allocation_list_default.append(n_final_strata)
                    allocation_dict_default[s['strata']] = int(n_final_strata)
                
                # Inisialisasi session state untuk alokasi jika belum ada
                if 'allocation_adjustments' not in st.session_state:
                    st.session_state['allocation_adjustments'] = allocation_list_default
                
                # Tampilkan tabel editor alokasi sampel per strata
                st.write("🎯 **Sesuaikan alokasi sampel per Strata (opsional)**:")
                
                allocation_editor_data = []
                for idx, s in enumerate(strata_summary):
                    allocation_editor_data.append({
                        "Strata": s['strata'],
                        "Populasi (N)": s['count'],
                        "Alokasi Otomatis": allocation_dict_default[s['strata']],
                        "Alokasi Manual": st.session_state['allocation_adjustments'][idx] 
                            if idx < len(st.session_state['allocation_adjustments']) 
                            else allocation_dict_default[s['strata']]
                    })
                
                # Data editor untuk customisasi alokasi
                edited_allocation_raw = st.data_editor(
                    allocation_editor_data,
                    column_config={
                        "Strata": "Range Nilai (Rupiah)",
                        "Populasi (N)": st.column_config.NumberColumn(
                            "Populasi",
                            disabled=True
                        ),
                        "Alokasi Otomatis": st.column_config.NumberColumn(
                            "Alokasi Otomatis",
                            disabled=True,
                            help="Berdasarkan rumus proporsi"
                        ),
                        "Alokasi Manual": st.column_config.NumberColumn(
                            "Alokasi Manual",
                            help="Sesuaikan jumlah sampel untuk strata ini",
                            required=True,
                            min_value=0
                        )
                    },
                    disabled=["Strata", "Populasi (N)", "Alokasi Otomatis"],
                    hide_index=True,
                    use_container_width=True
                )
                
                # Konversi hasil data_editor ke DataFrame
                edited_allocation = pd.DataFrame(edited_allocation_raw)
                
                # Proses hasil editan
                allocation_dict = {}
                final_allocations = []
                
                for idx, row in edited_allocation.iterrows():
                    strata_name = row['Strata']
                    n_adjusted = int(row['Alokasi Manual'])
                    pop_strata = row['Populasi (N)']
                    
                    # Validasi: jangan exceeding populasi
                    if n_adjusted > pop_strata:
                        n_adjusted = pop_strata
                        st.warning(f"⚠️ Alokasi untuk '{strata_name}' disesuaikan menjadi {pop_strata} (tidak boleh melebihi populasi)")
                    
                    allocation_dict[strata_name] = n_adjusted
                    final_allocations.append(n_adjusted)
                
                # Update n_res total dengan alokasi manual
                n_adjusted_total = sum(final_allocations)
                st.info(f"💡 Total Sampel (Sesuai Alokasi Manual): **{n_adjusted_total}** item.")
                n_res = n_adjusted_total
                
                # SIMPAN KE SESSION STATE
                st.session_state['allocation_dict'] = allocation_dict
                st.session_state['df_stratified'] = df
                st.session_state['allocation_adjustments'] = final_allocations

        except Exception as e:
            st.error(f"Gagal membagi kuartil otomatis. Data mungkin terlalu sedikit atau seragam. Error: {e}")
            n_res = 0

    # Tampilkan Hasil N
    if error_msg:
        st.error(error_msg)
        n_final = 0
    else:
        st.metric("Jumlah Sampel Disarankan (n)", f"{n_res} Item")
        n_final = st.number_input("Jumlah Sampel Final",
                                  min_value=1,
                                  max_value=len(df),
                                  value=int(n_res) if n_res > 0 else 30)

    # 4. TEKNIK EKSEKUSI
    st.markdown("---")
    st.subheader("Teknik Pemilihan Sampel")

    teknik = st.selectbox("Teknik Pemilihan", [
        "Acak Sederhana", "PPS (Wajib untuk MUS)", "Sistematis",
        "Sistematis Acak", "Stratifikasi (Top Value)",
        "Benford's Law (Anomali)"
    ])

    # --- PERBAIKAN: Gunakan Session State untuk tombol ---
    generate_btn = st.button("🚀 Generate Sampel")

    # Logika Hitung Sampel (Hanya jalan saat tombol diklik)
    if generate_btn:
        sampled_df = pd.DataFrame()

        # --- LOGIKA BARU: CEK APAKAH PAKAI STRATIFIED MPU ---
        if metode_sampling == "Stratified Mean Per Unit (MPU)" and 'allocation_dict' in st.session_state:
            st.info("Menggunakan pemilihan terdistribusi sesuai Strata...")
            
            # Ambil DF yang sudah ada label Strata-nya
            df_to_use = st.session_state.get('df_stratified', df)
            alloc_data = st.session_state['allocation_dict']

            # Panggil fungsi baru di selections.py
            sampled_df = sel.select_stratified_distributed(
                df_to_use,
                alloc_data,
                teknik, 
                value_col)

        else:
            # --- LOGIKA LAMA (UNTUK METODE NON-STRATIFIED) ---
            if teknik == "Acak Sederhana":
                sampled_df = sel.select_simple_random(df, n_final)
            elif teknik == "PPS (Wajib untuk MUS)":
                sampled_df = sel.select_pps(df, n_final, value_col)
            elif teknik == "Sistematis":
                sampled_df = sel.select_systematic(df, n_final)
            elif teknik == "Sistematis Acak":
                sampled_df = sel.select_random_systematic(df, n_final)
            elif teknik == "Stratifikasi (Top Value)":
                sampled_df = sel.select_stratified_top_value(df, n_final, value_col)
            elif teknik == "Benford's Law (Anomali)":
                sampled_df = sel.select_benford_anomaly(df, n_final, value_col)
        
        # SIMPAN HASIL KE SESSION STATE AGAR TIDAK HILANG SAAT REFRESH
        st.session_state['sampled_df'] = sampled_df
        # Reset state generate laporan agar tombol download docx hilang dulu jika generate ulang
        if 'report_docx_bytes' in st.session_state:
            del st.session_state['report_docx_bytes']

    # --- BAGIAN TAMPILAN (Dijalankan cek Session State, bukan tombol) ---
    if 'sampled_df' in st.session_state and not st.session_state['sampled_df'].empty:
        
        # Ambil data dari memory
        current_sampled_df = st.session_state['sampled_df']
        
        st.success(f"Terpilih {len(current_sampled_df)} sampel.")

        # Hitung total nilai sampel untuk kolom yang dipilih (jika ada)
        display_df = current_sampled_df.copy()
        try:
            if value_col in display_df.columns:
                total_sampled_value = display_df[value_col].sum()
                total_nilai_buku = df[value_col].sum()
                pct_nilai_sampel = (total_sampled_value / total_nilai_buku * 100) if total_nilai_buku > 0 else 0
                # Tampilkan metric ringkasan dalam 2 kolom
                col_metric1, col_metric2 = st.columns(2)
                with col_metric1:
                    st.metric("Total Nilai Sampel", f"Rp {total_sampled_value:,.2f}")
                with col_metric2:
                    st.metric("Persentase Nilai Sampel", f"{pct_nilai_sampel:.2f}%")
            else:
                st.info("Kolom Nilai Rupiah tidak ditemukan dalam data sampel; total tidak ditampilkan.")
        except Exception as e:
            st.warning(f"Gagal menghitung total nilai sampel: {e}")

        # Tampilkan rincian per strata (opsional)
        if 'Strata' in display_df.columns:
            st.write("Rincian Sampel per Strata:")
            st.write(display_df['Strata'].value_counts())

        st.dataframe(display_df)
        
        # === BUTTONS DOWNLOAD ===
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        with col_btn1:
            # Download Sampel Only
            buff_sampel = BytesIO()
            with pd.ExcelWriter(buff_sampel, engine='openpyxl') as writer:
                current_sampled_df.to_excel(writer, index=False, sheet_name='Sampel')
            buff_sampel.seek(0)
            st.download_button(
                label="📊 Download Sampel (.xlsx)",
                data=buff_sampel.getvalue(),
                file_name="sampel.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        with col_btn2:
            # Download Laporan Lengkap (.xlsx)
            buff_laporan = generate_laporan_xlsx(
                df,
                current_sampled_df,
                metode_sampling,
                teknik,
                confidence,
                sst,
                value_col,
                n_final
            )
            st.download_button(
                label="📋 Generate Laporan Lengkap (.xlsx)",
                data=buff_laporan.getvalue(),
                file_name=f"Laporan_Sampling_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        with col_btn3:
            # Generate & Download Laporan DOCX
            # Fungsi callback agar tidak refresh halaman total
            def on_generate_docx_inline():
                try:
                    # Ambil sampled_df dari session state
                    sdf = st.session_state['sampled_df']
                    buf = generate_laporan_docx(
                        df_original=df, # df diambil dari global scope (hasil upload)
                        sampled_df=sdf,
                        metode_sampling=metode_sampling,
                        teknik=teknik,
                        confidence=confidence,
                        sst=sst,
                        value_col=value_col,
                        n_final=n_final
                    )
                    st.session_state['report_docx_bytes'] = buf.getvalue()
                except Exception as exc:
                    st.error(f"❌ Gagal membuat .docx: {exc}")

            # Tombol pemicu generate docx
            st.button("📄 Generate Laporan (.docx)", on_click=on_generate_docx_inline, key="btn_docx_inline")
            
            # Tampilkan tombol download file-nya JIKA bytes sudah ada di session state
            if 'report_docx_bytes' in st.session_state and st.session_state['report_docx_bytes']:
                st.success("✅ Laporan .docx siap!")
                st.download_button(
                    label="⬇️ Download File (.docx)",
                    data=st.session_state['report_docx_bytes'],
                    file_name=f"Laporan_Sampling_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx_inline"
                )

    elif 'sampled_df' in st.session_state and st.session_state['sampled_df'].empty:
        st.warning("Tidak ada sampel yang terpilih. Cek parameter.")

else:
    st.info("Silakan upload data di sidebar.")


