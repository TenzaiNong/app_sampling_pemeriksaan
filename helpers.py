"""
Helper functions untuk aplikasi sampling audit
"""

import pandas as pd
import numpy as np
import csv
import re
import chardet
from io import BytesIO
from openpyxl.styles import Font, PatternFill, Alignment

import pendapatan_analyzer as pend_analyzer


def detect_encoding(uploaded_file):
    """
    Deteksi encoding file CSV.
    """
    try:
        uploaded_file.seek(0)
        raw_data = uploaded_file.read()
        
        result = chardet.detect(raw_data)
        encoding = result.get('encoding', 'utf-8')
        
        uploaded_file.seek(0)
        return encoding if encoding else 'utf-8'
    except Exception as e:
        uploaded_file.seek(0)
        return 'latin-1'


def detect_csv_delimiter(uploaded_file, sample_size=5000):
    """
    Deteksi delimiter CSV secara otomatis.
    """
    encoding = detect_encoding(uploaded_file)
    
    try:
        uploaded_file.seek(0)
        sample = uploaded_file.read(sample_size).decode(encoding, errors='ignore')
        
        delimiter = csv.Sniffer().sniff(sample).delimiter
        uploaded_file.seek(0)
        return delimiter, encoding
    except Exception as e:
        delimiters = [';', ',', '\t', '|']
        
        for delim in delimiters:
            try:
                uploaded_file.seek(0)
                test_df = pd.read_csv(uploaded_file, sep=delim, nrows=5, encoding=encoding)
                if len(test_df.columns) > 1:
                    uploaded_file.seek(0)
                    return delim, encoding
            except:
                continue
        
        uploaded_file.seek(0)
        return ',', encoding


def convert_rupiah_to_numeric(df):
    """
    Konversi kolom Rupiah ke numerik dengan perlindungan ketat.
    """
    print("--- Memulai Cek Konversi Data ---")

    for col in df.columns:
        try:
            if df[col].dtype != 'object':
                continue

            valid_samples = df[col].dropna().head(5).astype(str).tolist()
            
            if not valid_samples:
                continue

            is_text_column = False
            for s in valid_samples:
                clean_check = s.lower().replace('rp', '').replace('.', '').replace(',', '').replace('-', '').strip()
                
                if re.search('[a-z]', clean_check):
                    is_text_column = True
                    break
            
            if is_text_column:
                continue

            sample = valid_samples[0]
            
            if '.' in sample and ',' in sample:
                last_dot = sample.rfind('.')
                last_comma = sample.rfind(',')
                
                if last_dot < last_comma:
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace('Rp', '', regex=False)
                               .str.replace(' ', '', regex=False)
                               .str.replace('.', '', regex=False)
                               .str.replace(',', '.', regex=False)
                               .apply(pd.to_numeric, errors='coerce'))
                    print(f"✅ Konversi '{col}': Format Rupiah (Indo)")

            elif ',' in sample and '.' not in sample:
                clean_sample = sample.replace(',', '').strip()
                if clean_sample.isdigit():
                    df[col] = (df[col]
                               .astype(str)
                               .str.replace(',', '.', regex=False)
                               .apply(pd.to_numeric, errors='coerce'))
                    print(f"✅ Konversi '{col}': Desimal Koma")

        except Exception as e:
            pass
    
    print("--- Selesai Cek Konversi ---")
    return df


def generate_laporan_xlsx(df_original, sampled_df, metode_sampling, teknik, 
                          confidence, sst, value_col, n_final):
    """
    Generate laporan hasil sampling dalam format Excel (.xlsx)
    """
    from openpyxl import load_workbook
    
    buff = BytesIO()
    
    with pd.ExcelWriter(buff, engine='openpyxl') as writer:
        # === SHEET 1: RINGKASAN ===
        total_nilai_buku = df_original[value_col].sum()
        total_nilai_sampel = sampled_df[value_col].sum() if not sampled_df.empty else 0
        pct_nilai_sampel = (total_nilai_sampel / total_nilai_buku * 100) if total_nilai_buku > 0 else 0
        
        ringkasan_data = {
            'Keterangan': [
                'Total Populasi (N)',
                'Total Nilai Buku (Rp)',
                'Total Nilai Sampel (Rp)',
                'Persentase Nilai Sampel (%)',
                'Jumlah Sampel (n)',
                'Persentase Sampel (%)',
                'Metode Sampling',
                'Teknik Pemilihan',
                'Confidence Level (%)',
                'Salah Saji Tertoleransi (Rp)',
                'Tanggal Generate Laporan'
            ],
            'Nilai': [
                len(df_original),
                f"Rp {total_nilai_buku:,.2f}",
                f"Rp {total_nilai_sampel:,.2f}",
                f"{pct_nilai_sampel:.2f}%",
                n_final,
                f"{(n_final / len(df_original) * 100):.2f}%",
                metode_sampling,
                teknik,
                confidence,
                f"Rp {sst:,.2f}",
                pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S")
            ]
        }
        df_ringkasan = pd.DataFrame(ringkasan_data)
        df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
        
        # === SHEET 2: DETAIL SAMPEL ===
        sampled_df.to_excel(writer, sheet_name='Detail Sampel', index=False)
        
        # === SHEET 3: STATISTIK POPULASI ===
        statistik_data = {
            'Metrik': [
                'Mean',
                'Median',
                'Std Dev',
                'Min',
                'Max',
                'Q1 (25%)',
                'Q3 (75%)'
            ],
            'Populasi': [
                f"Rp {df_original[value_col].mean():,.2f}",
                f"Rp {df_original[value_col].median():,.2f}",
                f"Rp {df_original[value_col].std():,.2f}",
                f"Rp {df_original[value_col].min():,.2f}",
                f"Rp {df_original[value_col].max():,.2f}",
                f"Rp {df_original[value_col].quantile(0.25):,.2f}",
                f"Rp {df_original[value_col].quantile(0.75):,.2f}"
            ],
            'Sampel': [
                f"Rp {sampled_df[value_col].mean():,.2f}",
                f"Rp {sampled_df[value_col].median():,.2f}",
                f"Rp {sampled_df[value_col].std():,.2f}",
                f"Rp {sampled_df[value_col].min():,.2f}",
                f"Rp {sampled_df[value_col].max():,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.25):,.2f}",
                f"Rp {sampled_df[value_col].quantile(0.75):,.2f}"
            ]
        }
        df_statistik = pd.DataFrame(statistik_data)
        df_statistik.to_excel(writer, sheet_name='Statistik', index=False)
        
        # === FORMATTING ===
        workbook = writer.book
        
        ws_ringkasan = workbook['Ringkasan']
        fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        font_header = Font(bold=True, color="FFFFFF")
        
        for cell in ws_ringkasan[1]:
            cell.fill = fill_header
            cell.font = font_header
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        for column in ws_ringkasan.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws_ringkasan.column_dimensions[column_letter].width = adjusted_width
        
        ws_detail = workbook['Detail Sampel']
        for cell in ws_detail[1]:
            cell.fill = fill_header
            cell.font = font_header
        
        ws_statistik = workbook['Statistik']
        for cell in ws_statistik[1]:
            cell.fill = fill_header
            cell.font = font_header
    
    buff.seek(0)
    return buff


def generate_laporan_docx(df_original, sampled_df, metode_sampling, teknik,
                          confidence, sst, value_col, n_final, max_rows=300):
    """
    Generate laporan hasil sampling dalam format Word (.docx)
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    h = doc.add_heading('Laporan Sampling Pemeriksaan Keuangan', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tanggal: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S"))
    doc.add_paragraph(f"Metode Sampling: {metode_sampling}")
    doc.add_paragraph(f"Teknik Pemilihan: {teknik}")
    doc.add_paragraph(f"Confidence Level: {confidence}%")
    doc.add_paragraph(f"Salah Saji Tertoleransi (SST): Rp {sst:,.2f}")
    doc.add_paragraph(f"Jumlah Populasi (N): {len(df_original)}")
    if value_col:
        try:
            total_nilai_buku = df_original[value_col].sum()
            doc.add_paragraph(f"Total Nilai Buku: Rp {total_nilai_buku:,.2f}")
            total_nilai_sampel = sampled_df[value_col].sum() if not sampled_df.empty else 0
            pct_nilai = (total_nilai_sampel / total_nilai_buku * 100) if total_nilai_buku > 0 else 0
            doc.add_paragraph(f"Total Nilai Sampel: Rp {total_nilai_sampel:,.2f}")
            doc.add_paragraph(f"Persentase Nilai Sampel: {pct_nilai:.2f}%")
        except Exception:
            doc.add_paragraph("Total Nilai Buku: -")
    doc.add_paragraph(f"Jumlah Sampel (n): {n_final}")
    doc.add_paragraph()

    metrics = ['Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Q1 (25%)', 'Q3 (75%)']
    pop_values = []
    samp_values = []
    if value_col and value_col in df_original.columns and value_col in sampled_df.columns:
        pop = df_original[value_col].dropna().astype(float)
        samp = sampled_df[value_col].dropna().astype(float)
        pop_values = [
            f"Rp {pop.mean():,.2f}",
            f"Rp {pop.median():,.2f}",
            f"Rp {pop.std():,.2f}",
            f"Rp {pop.min():,.2f}",
            f"Rp {pop.max():,.2f}",
            f"Rp {pop.quantile(0.25):,.2f}",
            f"Rp {pop.quantile(0.75):,.2f}"
        ]
        samp_values = [
            f"Rp {samp.mean():,.2f}" if not samp.empty else "-",
            f"Rp {samp.median():,.2f}" if not samp.empty else "-",
            f"Rp {samp.std():,.2f}" if not samp.empty else "-",
            f"Rp {samp.min():,.2f}" if not samp.empty else "-",
            f"Rp {samp.max():,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.25):,.2f}" if not samp.empty else "-",
            f"Rp {samp.quantile(0.75):,.2f}" if not samp.empty else "-"
        ]
    else:
        pop_values = ['-'] * len(metrics)
        samp_values = ['-'] * len(metrics)

    doc.add_heading('Statistik Populasi vs Sampel', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = 'Metrik'
    hdr_cells[1].text = 'Populasi'
    hdr_cells[2].text = 'Sampel'
    for i, m in enumerate(metrics):
        row_cells = tbl.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = pop_values[i]
        row_cells[2].text = samp_values[i]

    doc.add_paragraph()

    doc.add_heading('Detail Sampel', level=2)
    if sampled_df.empty:
        doc.add_paragraph("Tidak ada sampel terpilih.")
    else:
        doc.add_paragraph(
            "Daftar sampel selengkapnya dilampirkan pada file Excel laporan. "
            "Untuk melihat data lengkap (semua kolom dan baris), silakan unduh laporan .xlsx."
        )

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- FUNGSI UNTUK TEMPLATE PENDAPATAN ---
def generate_template_pendapatan():
    """
    Generate template Excel untuk analisis Pendapatan WP.
    Kolom: Nomor, Nama WP, NPWPD, JANUARI-DESEMBER
    """
    from openpyxl import Workbook
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Template Pendapatan"
    
    # Header
    headers = ['NOMOR', 'NAMA WP', 'NPWPD'] + [
        'JANUARI', 'FEBRUARI', 'MARET', 'APRIL', 'MEI', 'JUNI',
        'JULI', 'AGUSTUS', 'SEPTEMBER', 'OKTOBER', 'NOVEMBER', 'DESEMBER'
    ]
    
    ws.append(headers)
    
    # Format header
    fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    font_header = Font(bold=True, color="FFFFFF", size=11)
    alignment_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    alignment_right = Alignment(horizontal="right", vertical="center")
    
    for cell in ws[1]:
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = alignment_center
    
    # Tambah contoh data (opsional)
    ws.append([1, 'Contoh WP 1', '12.345.678.901-123.456', 1000000, 1000000, 1000000, 1000000, 
               1000000, 1000000, 1000000, 1000000, 1000000, 1000000, 1000000, 1000000])
    
    # Set lebar kolom
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 20
    
    for col in ['D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O']:
        ws.column_dimensions[col].width = 14
        # Format nilai sebagai currency
        for row in ws.iter_rows(min_col=ws[col+'1'].column, max_col=ws[col+'1'].column, min_row=2):
            for cell in row:
                cell.alignment = alignment_right
    
    buff = BytesIO()
    wb.save(buff)
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK LAPORAN PENDAPATAN (EXCEL) ---
def generate_laporan_pendapatan_xlsx(df_original, anomali_list, bulan_cols):
    """
    Generate laporan analisis anomali Pendapatan dalam format Excel.
    """
    from openpyxl import Workbook
    
    buff = BytesIO()
    
    # Hitung statistik
    statistik = pend_analyzer.hitung_statistik_pendapatan(df_original, bulan_cols)
    
    with pd.ExcelWriter(buff, engine='openpyxl') as writer:
        # === SHEET 1: RINGKASAN ===
        # Hitung total realisasi untuk WP anomali
        anomalous_total = sum([item.get('total_realisasi', 0) for item in anomali_list]) if anomali_list else 0
        total_pendapatan = statistik.get('total_pendapatan', 0)
        anomalous_pct = (anomalous_total / total_pendapatan * 100) if total_pendapatan > 0 else 0

        ringkasan_data = {
            'Keterangan': [
                'Total WP dalam Populasi',
                'WP dengan Anomali Terdeteksi',
                'Persentase WP Anomali (%)',
                'Total Realisasi Anomali (Rp)',
                'Persentase Realisasi Anomali (%)',
                'Total Pendapatan (Rp)',
                'Rata-rata Pendapatan (Rp)',
                'Median Pendapatan (Rp)',
                'Standar Deviasi (Rp)',
                'Pendapatan Minimum (Rp)',
                'Pendapatan Maksimum (Rp)',
                'Tanggal Generate Laporan'
            ],
            'Nilai': [
                statistik['total_wp'],
                len(anomali_list),
                f"{(len(anomali_list) / statistik['total_wp'] * 100):.2f}%" if statistik['total_wp'] > 0 else "0%",
                f"Rp {anomalous_total:,.2f}",
                f"{anomalous_pct:.2f}%",
                f"Rp {statistik['total_pendapatan']:,.2f}",
                f"Rp {statistik['rata_rata']:,.2f}",
                f"Rp {statistik['median']:,.2f}",
                f"Rp {statistik['std_dev']:,.2f}",
                f"Rp {statistik['min']:,.2f}",
                f"Rp {statistik['max']:,.2f}",
                pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S")
            ]
        }
        df_ringkasan = pd.DataFrame(ringkasan_data)
        df_ringkasan.to_excel(writer, sheet_name='Ringkasan', index=False)
        
        # === SHEET 2: DAFTAR ANOMALI ===
        if anomali_list:
            df_anomali = pd.DataFrame(anomali_list)
            # Rename kolom untuk tampilan yang lebih baik
            df_anomali = df_anomali.rename(columns={
                'nomor': 'No',
                'nama_wp': 'Nama WP',
                'npwpd': 'NPWPD',
                'jenis_anomali': 'Jenis Anomali',
                'bulan_terisi': 'Bulan Terisi',
                'rata_rata': 'Rata-rata (Rp)',
                'min': 'Min (Rp)',
                'max': 'Max (Rp)',
                'std_dev': 'Std Dev (Rp)'
            })
            # Jika analyzer mengembalikan total_realisasi, tambahkan kolom untuk ditampilkan
            if 'total_realisasi' in df_anomali.columns:
                df_anomali = df_anomali.rename(columns={'total_realisasi': 'Total Realisasi (Rp)'})
            df_anomali.to_excel(writer, sheet_name='Daftar Anomali', index=False)
        
        # === FORMATTING ===
        workbook = writer.book
        
        # Format Ringkasan
        ws_ringkasan = workbook['Ringkasan']
        fill_header = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        font_header = Font(bold=True, color="FFFFFF")
        
        for cell in ws_ringkasan[1]:
            cell.fill = fill_header
            cell.font = font_header
        
        for column in ws_ringkasan.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws_ringkasan.column_dimensions[column_letter].width = (max_length + 2)
        
        # Format Daftar Anomali
        if 'Daftar Anomali' in workbook.sheetnames:
            ws_anomali = workbook['Daftar Anomali']
            for cell in ws_anomali[1]:
                cell.fill = fill_header
                cell.font = font_header
    
    buff.seek(0)
    return buff


# --- FUNGSI UNTUK LAPORAN PENDAPATAN (DOCX) ---
def generate_laporan_pendapatan_docx(df_original, anomali_list, bulan_cols):
    """
    Generate laporan analisis anomali Pendapatan dalam format Word.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # Judul
    h = doc.add_heading('Laporan Analisis Anomali Pendapatan Wajib Pajak', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tanggal Generate: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%d-%m-%Y %H:%M:%S"))
    
    # Statistik Umum
    statistik = pend_analyzer.hitung_statistik_pendapatan(df_original, bulan_cols)
    
    doc.add_heading('Ringkasan Hasil Analisis', level=2)
    doc.add_paragraph(f"Total WP dalam Populasi: {statistik['total_wp']} WP")
    doc.add_paragraph(f"WP dengan Anomali Terdeteksi: {len(anomali_list)} WP")
    doc.add_paragraph(f"Persentase WP Anomali: {(len(anomali_list) / statistik['total_wp'] * 100):.2f}%")
    # Hitung total realisasi anomali dan persentasenya
    anomalous_total = sum([item.get('total_realisasi', 0) for item in anomali_list]) if anomali_list else 0
    total_pendapatan = statistik.get('total_pendapatan', 0)
    anomalous_pct = (anomalous_total / total_pendapatan * 100) if total_pendapatan > 0 else 0
    doc.add_paragraph(f"Total Realisasi Anomali: Rp {anomalous_total:,.2f}")
    doc.add_paragraph(f"Persentase Realisasi Anomali: {anomalous_pct:.2f}%")
    
    doc.add_heading('Statistik Pendapatan', level=2)
    doc.add_paragraph(f"Total Pendapatan: Rp {statistik['total_pendapatan']:,.2f}")
    doc.add_paragraph(f"Rata-rata Pendapatan: Rp {statistik['rata_rata']:,.2f}")
    doc.add_paragraph(f"Median Pendapatan: Rp {statistik['median']:,.2f}")
    doc.add_paragraph(f"Standar Deviasi: Rp {statistik['std_dev']:,.2f}")
    doc.add_paragraph(f"Pendapatan Minimum: Rp {statistik['min']:,.2f}")
    doc.add_paragraph(f"Pendapatan Maksimum: Rp {statistik['max']:,.2f}")
    
    # Daftar WP dengan Anomali
    doc.add_heading('Daftar WP dengan Anomali', level=2)
    
    if anomali_list:
        # Buat tabel (tambah kolom Total Realisasi)
        tbl = doc.add_table(rows=1, cols=6)
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Nama WP'
        hdr_cells[2].text = 'NPWPD'
        hdr_cells[3].text = 'Jenis Anomali'
        hdr_cells[4].text = 'Rata-rata Pendapatan'
        hdr_cells[5].text = 'Total Realisasi'

        for item in anomali_list[:100]:  # Limit 100 baris untuk dokumen
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(item.get('nomor', ''))
            row_cells[1].text = str(item.get('nama_wp', ''))
            row_cells[2].text = str(item.get('npwpd', ''))
            row_cells[3].text = str(item.get('jenis_anomali', ''))
            row_cells[4].text = f"Rp {item.get('rata_rata', 0):,.2f}"
            row_cells[5].text = f"Rp {item.get('total_realisasi', 0):,.2f}"
        
        if len(anomali_list) > 100:
            doc.add_paragraph(f"... dan {len(anomali_list) - 100} WP lainnya (lihat laporan Excel untuk lengkapnya)")
    else:
        doc.add_paragraph("Tidak ditemukan WP dengan anomali berdasarkan kriteria analisis.")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
